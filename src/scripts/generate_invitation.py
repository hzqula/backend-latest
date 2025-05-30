# backend/scripts/generate_invitation.py
from docx import Document
import sys
import json
import os
from docx2pdf import convert

def replace_in_paragraphs(paragraphs, placeholder, replacement):
    for paragraph in paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, replacement)
            print(f"Replaced {placeholder} with {replacement} in paragraph")

def replace_in_tables(tables, placeholder, replacement):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = cell.text.replace(placeholder, replacement)
                    print(f"Replaced {placeholder} with {replacement} in table cell")

def generate_seminar_invitation(seminar_data, template_path, output_path):
    print(f"Received seminar data: {seminar_data}")
    print(f"Loading template from: {template_path}")
    if not os.path.exists(template_path):
        print(f"Error: Template file not found at {template_path}", file=sys.stderr)
        sys.exit(1)

    doc = Document(template_path)

    # Daftar penggantian
    replacements = [
        ("<<SEMINAR>>", "Seminar Proposal"),
        ("<<KETUA SEMINAR>>", seminar_data.get("ketua_seminar", "Tidak ditentukan")),
        ("<<NAMA PEMBIMBING 1>>", seminar_data.get("pembimbing_1", "Tidak ditentukan")),
        ("<<NAMA PEMBIMBING 2>>", seminar_data.get("pembimbing_2", "Tidak ditentukan")),
        ("<<NAMA PENGUJI 1>>", seminar_data.get("penguji_1", "Tidak ditentukan")),
        ("<<NAMA PENGUJI 2>>", seminar_data.get("penguji_2", "Tidak ditentukan")),
        ("<<NAMA MAHASISWA>>", seminar_data.get("student_name", "Tidak diketahui")),
        ("<<NIM>>", seminar_data.get("nim", "Tidak diketahui")),
        ("<<JUDUL PENELITIAN>>", seminar_data.get("judul_penelitian", "Tidak diketahui")),
        ("<<HARI>>", seminar_data.get("hari", "Tidak diketahui")),
        ("<<TANGGAL SEMINAR>>", seminar_data.get("date", "Tidak diketahui")),
        ("<<WAKTU SEMINAR>>", seminar_data.get("time", "Tidak diketahui")),
        ("<<RUANG SEMINAR>>", seminar_data.get("room", "Tidak ditentukan")),
    ]

    # Ganti di paragraf
    for placeholder, replacement in replacements:
        replace_in_paragraphs(doc.paragraphs, placeholder, replacement)

    # Ganti di tabel
    for placeholder, replacement in replacements:
        replace_in_tables(doc.tables, placeholder, replacement)

    # Ganti di "Pekanbaru, <<TANGGAL SEMINAR>>"
    for paragraph in doc.paragraphs:
        if "Pekanbaru, <<TANGGAL SEMINAR>>" in paragraph.text:
            paragraph.text = paragraph.text.replace("<<TANGGAL SEMINAR>>", seminar_data.get("date", "Tidak diketahui"))
            print(f"Replaced Pekanbaru, <<TANGGAL SEMINAR>> with Pekanbaru, {seminar_data.get('date', 'Tidak diketahui')}")

    # Simpan sebagai .docx sementara
    docx_output_path = output_path
    print(f"Saving DOCX to: {docx_output_path}")
    doc.save(docx_output_path)
    print(f"DOCX saved successfully: {docx_output_path}")

    # Konversi .docx ke .pdf
    pdf_output_path = output_path.replace(".docx", ".pdf")
    print(f"Converting DOCX to PDF: {pdf_output_path}")
    convert(docx_output_path, pdf_output_path)
    print(f"PDF saved successfully: {pdf_output_path}")

    # Hapus file .docx sementara
    os.remove(docx_output_path)
    print(f"Temporary DOCX file removed: {docx_output_path}")

    # Update output_path menjadi path file PDF
    seminar_data["output_path"] = pdf_output_path

if __name__ == "__main__":
    try:
        seminar_data = json.loads(sys.argv[1])
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        template_path = os.path.join(base_dir, "templates/undangan_seminar.docx")
        output_path = seminar_data["output_path"]
        generate_seminar_invitation(seminar_data, template_path, output_path)
    except Exception as e:
        print(f"Error in Python script: {str(e)}", file=sys.stderr)
        sys.exit(1)